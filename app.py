import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import datetime
import json
import os
import sv_ttk
from scheduler import generate_daily_schedule, parse_dt, calculate_points

class App:
    def __init__(self, root):
        self.root = root
        self.root.title("Gestor Contínuo de Escalas v1.2")
        self.root.geometry("1100x800")
        
        self.dispensas = {} # {pessoa_id: [(start_date, end_date)]}
        self.schedule_result = None
        
        import sys
        if getattr(sys, 'frozen', False):
            base_dir = os.path.dirname(sys.executable)
        else:
            base_dir = os.path.dirname(os.path.abspath(__file__))
            
        self.state_file = os.path.join(base_dir, "state.json")
        
        # Definir Ícone do App
        try:
            icon_path = self.get_resource_path("brasao.png")
            if os.path.exists(icon_path):
                img_icon = tk.PhotoImage(file=icon_path)
                self.root.iconphoto(False, img_icon)
        except Exception as e:
            print("Erro ao carregar ícone:", e)
        
        self.UNIDADES_DATA = {
            "BC": {"nome": "BATERIA DE COMANDO", "sigla": "Bia C", "sigla_doc": "BC", "id_start": 301},
            "1BO": {"nome": "1ª BATERIA DE OBUSES", "sigla": "1ª BO", "sigla_doc": "1ª BO", "id_start": 401},
            "2BO": {"nome": "2ª BATERIA DE OBUSES", "sigla": "2ª BO", "sigla_doc": "2ª BO", "id_start": 501},
        }
        
        self.current_state = self.load_state()
        self.preview_state = None
        
        self.dispensas = self._parse_dispensas_from_state()
        
        self.setup_ui()
        self.atualizar_lista_pessoas()
        self.atualizar_lista_dispensas()
        self.atualizar_data_alvo()
        self.atualizar_lista_historico()
        self.atualizar_ranking()
        
    def solicitar_unidade(self):
        selection = {"val": "BC"}
        win = tk.Toplevel(self.root)
        win.title("Configuração Inicial")
        win.geometry("400x300")
        win.transient(self.root)
        win.grab_set()
        
        ttk.Label(win, text="Selecione a Subunidade:", font=("Segoe UI", 12, "bold")).pack(pady=20)
        
        var = tk.StringVar(value="BC")
        ttk.Radiobutton(win, text="Bateria de Comando (Soldados 301+)", variable=var, value="BC").pack(anchor=tk.W, padx=50, pady=5)
        ttk.Radiobutton(win, text="1ª Bateria de Obuses (Soldados 401+)", variable=var, value="1BO").pack(anchor=tk.W, padx=50, pady=5)
        ttk.Radiobutton(win, text="2ª Bateria de Obuses (Soldados 501+)", variable=var, value="2BO").pack(anchor=tk.W, padx=50, pady=5)
        
        def confirmar():
            selection["val"] = var.get()
            win.destroy()
            
        ttk.Button(win, text="Confirmar e Iniciar", command=confirmar, style="Accent.TButton").pack(pady=30)
        
        self.root.wait_window(win)
        return selection["val"]

    def load_state(self):
        if os.path.exists(self.state_file):
            try:
                with open(self.state_file, 'r', encoding='utf-8') as f:
                    state = json.load(f)
                    if 'unidade' not in state:
                        state['unidade'] = self.solicitar_unidade()
                        self.save_state(state)
                    return state
            except Exception as e:
                print("Erro ao carregar state.json:", e)
        
        # Novo Estado Inicial
        unidade = self.solicitar_unidade()
        id_start = self.UNIDADES_DATA[unidade]["id_start"]
        
        default_state = {
            'unidade': unidade,
            'pessoas': {}, 
            'historico_escalas': [], 
            'dispensas': {},
            'nome_cmt': "RENAN LOUREIRO LENTZ - Cap",
            'nome_sgte': "HEBERT CARLOS VIANA - 2° Sgt"
        }
        for i in range(id_start, id_start + 60):
            default_state['pessoas'][str(i)] = {"ativo": True}
            
        return default_state
        
    def save_state(self, state):
        try:
            with open(self.state_file, 'w', encoding='utf-8') as f:
                json.dump(state, f, ensure_ascii=False, indent=4)
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao salvar histórico:\n{e}")



    def get_next_date(self):
        historico = self.current_state.get('historico_escalas', [])
        if historico:
            # Pega a última data registrada no histórico
            last_date_str = historico[-1]['data']
            last_date = parse_dt(last_date_str)
            if last_date:
                return last_date + datetime.timedelta(days=1)
        return datetime.date.today() + datetime.timedelta(days=1)

    def setup_ui(self):
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        title_lbl = ttk.Label(main_frame, text="Gestão Contínua de Escalas", font=("Segoe UI", 24, "bold"))
        title_lbl.pack(pady=(0, 15))
        
        # Tabs
        self.notebook = ttk.Notebook(main_frame)
        self.notebook.pack(fill=tk.BOTH, expand=True)
        
        self.tab_efetivo = ttk.Frame(self.notebook, padding=10)
        self.tab_gerador = ttk.Frame(self.notebook, padding=10)
        self.tab_historico = ttk.Frame(self.notebook, padding=10)
        self.tab_ranking = ttk.Frame(self.notebook, padding=10)
        self.tab_config = ttk.Frame(self.notebook, padding=10)
        
        self.notebook.add(self.tab_gerador, text="1. Gerar Escala Diária")
        self.notebook.add(self.tab_historico, text="2. Histórico")
        self.notebook.add(self.tab_efetivo, text="3. Efetivo e Dispensas")
        self.notebook.add(self.tab_ranking, text="4. Rank de Cansaço")
        self.notebook.add(self.tab_config, text="5. Configurações")
        
        self.setup_tab_efetivo()
        self.setup_tab_gerador()
        self.setup_tab_historico()
        self.setup_tab_ranking()
        self.setup_tab_config()

    def setup_tab_efetivo(self):
        paned = ttk.PanedWindow(self.tab_efetivo, orient=tk.HORIZONTAL)
        paned.pack(fill=tk.BOTH, expand=True)
        
        # Cadastro
        cadastro_frame = ttk.LabelFrame(paned, text="Cadastro de Efetivo", padding="10")
        paned.add(cadastro_frame, weight=1)
        
        add_frame = ttk.Frame(cadastro_frame)
        add_frame.pack(fill=tk.X, pady=5)
        ttk.Label(add_frame, text="Nº:").pack(side=tk.LEFT)
        self.entry_add_pessoa = ttk.Entry(add_frame, width=10)
        self.entry_add_pessoa.pack(side=tk.LEFT, padx=5)
        ttk.Button(add_frame, text="Cadastrar", command=self.cadastrar_pessoa).pack(side=tk.LEFT)
        
        self.listbox_pessoas = tk.Listbox(cadastro_frame, height=15)
        self.listbox_pessoas.pack(fill=tk.BOTH, expand=True, pady=5)
        
        btn_frame = ttk.Frame(cadastro_frame)
        btn_frame.pack(fill=tk.X)
        ttk.Button(btn_frame, text="Alternar Ativo/Inativo (Rota)", command=self.alternar_status_pessoa).pack(side=tk.LEFT, fill=tk.X, expand=True, padx=2)
        ttk.Button(btn_frame, text="Excluir", command=self.excluir_pessoa).pack(side=tk.LEFT, fill=tk.X, expand=True, padx=2)

        btn_frame_2 = ttk.Frame(cadastro_frame)
        btn_frame_2.pack(fill=tk.X, pady=2)
        ttk.Button(btn_frame_2, text="Alternar PO", command=self.alternar_po_pessoa).pack(side=tk.LEFT, fill=tk.X, expand=True, padx=2)
        ttk.Button(btn_frame_2, text="Alternar SGT", command=self.alternar_sargentiacao_pessoa).pack(side=tk.LEFT, fill=tk.X, expand=True, padx=2)

        # Dispensas
        dispensas_frame = ttk.LabelFrame(paned, text="Dispensas Médicas / Férias", padding="10")
        paned.add(dispensas_frame, weight=1)
        
        d_top = ttk.Frame(dispensas_frame)
        d_top.pack(fill=tk.X)
        ttk.Label(d_top, text="Nº:").grid(row=0, column=0)
        self.entry_disp_pessoa = ttk.Entry(d_top, width=6)
        self.entry_disp_pessoa.grid(row=0, column=1)
        ttk.Label(d_top, text="Início:").grid(row=0, column=2)
        self.entry_disp_inicio = ttk.Entry(d_top, width=10)
        self.entry_disp_inicio.grid(row=0, column=3)
        ttk.Label(d_top, text="Fim:").grid(row=1, column=0)
        self.entry_disp_fim = ttk.Entry(d_top, width=10)
        self.entry_disp_fim.grid(row=1, column=1, columnspan=2)
        ttk.Button(d_top, text="Add", command=self.adicionar_dispensa).grid(row=1, column=3)
        
        self.listbox_dispensas = tk.Listbox(dispensas_frame, height=10)
        self.listbox_dispensas.pack(fill=tk.BOTH, expand=True, pady=5)
        ttk.Button(dispensas_frame, text="Remover Dispensa", command=self.remover_dispensa).pack()

    def setup_tab_gerador(self):
        gerador_frame = ttk.LabelFrame(self.tab_gerador, text="Geração de Escala", padding="15")
        gerador_frame.pack(fill=tk.X, pady=5)
        
        self.lbl_target_date = ttk.Label(gerador_frame, text="Data Alvo: --/--/----", font=("Segoe UI", 16, "bold"))
        self.lbl_target_date.pack(pady=5)
        
        opts_frame = ttk.Frame(gerador_frame)
        opts_frame.pack(pady=10)
        
        self.var_guarda = tk.BooleanVar(value=False)
        self.var_plantao = tk.BooleanVar(value=True)
        self.var_apoio = tk.BooleanVar(value=False)
        
        ttk.Checkbutton(opts_frame, text="Incluir GUARDA (24)", variable=self.var_guarda).pack(side=tk.LEFT, padx=10)
        ttk.Checkbutton(opts_frame, text="Incluir PLANTÃO (6)", variable=self.var_plantao).pack(side=tk.LEFT, padx=10)
        ttk.Checkbutton(opts_frame, text="Incluir APOIO (2)", variable=self.var_apoio).pack(side=tk.LEFT, padx=10)
        
        self.manual_gc_frame = ttk.LabelFrame(gerador_frame, text="Preenchimento Manual de Funções (Guarnição)", padding="10")
        self.manual_gc_frame.pack(fill=tk.X, pady=10)
        
        self.manual_entries = {}
        self.manual_categories = {}
        self.funcoes_list = [
            ("OF DIA", "of_dia"), ("ADJ OF DIA", "adj_of_dia"), 
            ("SGT DIA BIA C", "sgt_dia_bia_c"), ("CB DIA BIA C", "cb_dia_bia_c"), 
            ("MOT DIA", "mot_dia"), ("PADIOLEIRO", "padioleiro"), 
            ("SOMBRA", "sombra"), ("MOT VILA", "mot_vila"), ("GDA VILA", "gda_vila"),
            ("CMT GDA", "cmt_gda"), ("CMT GDA VILA", "cmt_gda_vila"), ("SGT DIA BIA O", "sgt_dia_bia_o"),
            ("CB GDA QTEL", "cb_gda_qtel"), ("CB GDA VILA", "cb_gda_vila"), ("CB DIA BIA O", "cb_dia_bia_o"),
            ("MOT SUP DIA", "mot_sup_dia"), ("GDA QTEL EP", "gda_qtel_ep"), ("REFORÇO EP", "reforco_ep"),
            ("PERM. HT", "permanencia_ht"), ("REFORÇO EV", "reforco_ev")
        ]
        
        has_cat = ["mot_dia", "padioleiro", "mot_vila"]
        
        for i, (label, key) in enumerate(self.funcoes_list):
            row = i // 4
            col = (i % 4) * 2
            
            f_frame = ttk.Frame(self.manual_gc_frame)
            f_frame.grid(row=row, column=col, columnspan=2, padx=2, pady=5, sticky=tk.W)
            
            ttk.Label(f_frame, text=label+":").pack(side=tk.LEFT)
            entry = ttk.Entry(f_frame, width=12)
            entry.pack(side=tk.LEFT, padx=2)
            self.manual_entries[key] = entry
            
            if key in has_cat:
                cat_var = tk.StringVar(value="SD EP")
                cb = ttk.Combobox(f_frame, textvariable=cat_var, values=["CB", "SD EP", "SD EV"], width=6, state="readonly")
                cb.pack(side=tk.LEFT)
                self.manual_categories[key] = cb

        actions_frame = ttk.Frame(gerador_frame)
        actions_frame.pack(fill=tk.X, pady=10)
        
        ttk.Button(actions_frame, text="1. GERAR PRÉVIA DO DIA", command=self.gerar_escala, style="Accent.TButton").pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        self.btn_confirm = ttk.Button(actions_frame, text="2. CONFIRMAR E SALVAR", command=self.confirmar_escala, state=tk.DISABLED)
        self.btn_confirm.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        self.btn_print = ttk.Button(actions_frame, text="IMPRIMIR DIA", command=self.imprimir_escala, state=tk.DISABLED)
        self.btn_print.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)

        details_frame = ttk.LabelFrame(self.tab_gerador, text="Visualização da Prévia", padding="10")
        details_frame.pack(fill=tk.BOTH, expand=True, pady=10)
        
        self.text_details = tk.Text(details_frame, wrap=tk.WORD, font=("Consolas", 14))
        self.text_details.pack(fill=tk.BOTH, expand=True)

    def setup_tab_historico(self):
        paned = ttk.PanedWindow(self.tab_historico, orient=tk.HORIZONTAL)
        paned.pack(fill=tk.BOTH, expand=True)
        
        list_frame = ttk.LabelFrame(paned, text="Escalas Salvas", padding="10")
        paned.add(list_frame, weight=1)
        
        self.listbox_historico = tk.Listbox(list_frame, height=15, font=("Consolas", 12))
        self.listbox_historico.pack(fill=tk.BOTH, expand=True, pady=5)
        self.listbox_historico.bind('<<ListboxSelect>>', self.on_historico_select)
        
        btn_frame = ttk.Frame(list_frame)
        btn_frame.pack(fill=tk.X)
        ttk.Button(btn_frame, text="Excluir Escala Selecionada", command=self.excluir_historico).pack(side=tk.LEFT, fill=tk.X, expand=True, padx=2)
        ttk.Button(btn_frame, text="Imprimir Selecionada", command=self.imprimir_historico).pack(side=tk.LEFT, fill=tk.X, expand=True, padx=2)
        
        view_frame = ttk.LabelFrame(paned, text="Detalhes da Escala", padding="10")
        paned.add(view_frame, weight=2)
        
        self.text_hist_details = tk.Text(view_frame, wrap=tk.WORD, font=("Consolas", 14))
        self.text_hist_details.pack(fill=tk.BOTH, expand=True)

    def setup_tab_ranking(self):
        ranking_frame = ttk.Frame(self.tab_ranking, padding="10")
        ranking_frame.pack(fill=tk.BOTH, expand=True)
        
        ttk.Label(ranking_frame, text="Ranking de Cansaço (Pontuação Acumulada)", font=("Segoe UI", 16, "bold")).pack(pady=(0, 10))
        
        cols = ("Soldado", "Pontos Preta", "Pontos Vermelha", "Total")
        self.tree_ranking = ttk.Treeview(ranking_frame, columns=cols, show='headings')
        
        for col in cols:
            self.tree_ranking.heading(col, text=col)
            self.tree_ranking.column(col, width=150, anchor=tk.CENTER)
            
        self.tree_ranking.pack(fill=tk.BOTH, expand=True)
        
        ttk.Button(ranking_frame, text="Atualizar Ranking", command=self.atualizar_ranking).pack(pady=10)

    def atualizar_ranking(self):
        for item in self.tree_ranking.get_children():
            self.tree_ranking.delete(item)
            
        pessoas_db = self.current_state.get('pessoas', {})
        historico = self.current_state.get('historico_escalas', [])
        
        # Calcula pontos dinamicamente do histórico
        p_preta, p_vermelha, _, _, _ = calculate_points(historico, list(pessoas_db.keys()))
        
        ranking_list = []
        for p in pessoas_db:
            preta = p_preta.get(p, 0)
            vermelha = p_vermelha.get(p, 0)
            total = preta + vermelha
            ranking_list.append((p, preta, vermelha, total))
            
        # Ordena pelo total (mais cansados primeiro)
        ranking_list.sort(key=lambda x: x[3], reverse=True)
        
        for entry in ranking_list:
            self.tree_ranking.insert('', tk.END, values=entry)

    def setup_tab_config(self):
        config_frame = ttk.LabelFrame(self.tab_config, text="Configurações da Unidade", padding="20")
        config_frame.pack(fill=tk.X, pady=10)
        
        ttk.Label(config_frame, text="Nome do Comandante (Capitão):", font=("Segoe UI", 12)).pack(anchor=tk.W, pady=(0, 5))
        self.entry_cmt = ttk.Entry(config_frame, font=("Segoe UI", 12))
        self.entry_cmt.pack(fill=tk.X, pady=(0, 15))
        self.entry_cmt.insert(0, self.current_state.get('nome_cmt', "RENAN LOUREIRO LENTZ - Cap"))

        ttk.Label(config_frame, text="Subunidade Atual:", font=("Segoe UI", 12)).pack(anchor=tk.W, pady=(0, 5))
        self.combo_unidade = ttk.Combobox(config_frame, values=["BC", "1BO", "2BO"], font=("Segoe UI", 12), state="readonly")
        self.combo_unidade.pack(fill=tk.X, pady=(0, 15))
        self.combo_unidade.set(self.current_state.get('unidade', "BC"))
        
        ttk.Label(config_frame, text="Nome do Sargenteante:", font=("Segoe UI", 12)).pack(anchor=tk.W, pady=(0, 5))
        self.entry_sgte = ttk.Entry(config_frame, font=("Segoe UI", 12))
        self.entry_sgte.pack(fill=tk.X, pady=(0, 15))
        self.entry_sgte.insert(0, self.current_state.get('nome_sgte', "HEBERT CARLOS VIANA - 2° Sgt"))
        
        # Novos campos para números de Aditamento e Boletim
        nums_frame = ttk.Frame(config_frame)
        nums_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(nums_frame, text="Nr Aditamento:", font=("Segoe UI", 11)).grid(row=0, column=0, sticky=tk.W, padx=5)
        self.entry_nr_adit = ttk.Entry(nums_frame, width=10, font=("Segoe UI", 11))
        self.entry_nr_adit.grid(row=0, column=1, sticky=tk.W, padx=5)
        self.entry_nr_adit.insert(0, self.current_state.get('nr_aditamento', "___"))
        
        ttk.Label(nums_frame, text="Nr Boletim Interno:", font=("Segoe UI", 11)).grid(row=0, column=2, sticky=tk.W, padx=5)
        self.entry_nr_bol = ttk.Entry(nums_frame, width=10, font=("Segoe UI", 11))
        self.entry_nr_bol.grid(row=0, column=3, sticky=tk.W, padx=5)
        self.entry_nr_bol.insert(0, self.current_state.get('nr_boletim', "___"))
        
        ttk.Button(config_frame, text="Salvar Configurações", command=self.salvar_config, style="Accent.TButton").pack(fill=tk.X, pady=10)

    def salvar_config(self):
        self.current_state['nome_cmt'] = self.entry_cmt.get().strip()
        self.current_state['nome_sgte'] = self.entry_sgte.get().strip()
        self.current_state['nr_aditamento'] = self.entry_nr_adit.get().strip()
        self.current_state['nr_boletim'] = self.entry_nr_bol.get().strip()
        self.current_state['unidade'] = self.combo_unidade.get()
        self.save_state(self.current_state)
        messagebox.showinfo("Sucesso", "Configurações salvas com sucesso!")

    def parse_dt_local(self, date_str):
        try: return datetime.datetime.strptime(date_str, "%d/%m/%Y").date()
        except: return None
        
    def get_resource_path(self, relative_path):
        import sys
        try:
            base_path = sys._MEIPASS
        except Exception:
            base_path = os.path.abspath(".")
        return os.path.join(base_path, relative_path)

    def _parse_dispensas_from_state(self):
        dispensas_str = self.current_state.get('dispensas', {})
        disp_parsed = {}
        for p_str, dates in dispensas_str.items():
            p_int = int(p_str)
            disp_parsed[p_int] = []
            for d_start_str, d_end_str in dates:
                try:
                    d_start = datetime.datetime.strptime(d_start_str, "%Y-%m-%d").date()
                    d_end = datetime.datetime.strptime(d_end_str, "%Y-%m-%d").date()
                    disp_parsed[p_int].append((d_start, d_end))
                except: pass
        return disp_parsed
        
    def _sync_dispensas_to_state(self):
        dispensas_str = {}
        for p_int, dates in self.dispensas.items():
            dispensas_str[str(p_int)] = []
            for d_start, d_end in dates:
                dispensas_str[str(p_int)].append((d_start.strftime("%Y-%m-%d"), d_end.strftime("%Y-%m-%d")))
        self.current_state['dispensas'] = dispensas_str

    # --- EFETIVO LOGIC ---
    def atualizar_lista_pessoas(self):
        self.listbox_pessoas.delete(0, tk.END)
        pessoas_db = self.current_state.get('pessoas', {})
        sorted_keys = sorted(pessoas_db.keys(), key=lambda x: int(x) if x.isdigit() else x)
        for k in sorted_keys:
            ativo = pessoas_db[k].get('ativo', True)
            status = "ATIVO" if ativo else "INATIVO (ROTA)"
            
            tags = []
            if pessoas_db[k].get('is_po', False):
                tags.append("[PO]")
            if pessoas_db[k].get('is_sargentiacao', False):
                tags.append("[SGT]")
            
            tag_str = (" " + " ".join(tags)) if tags else ""
            self.listbox_pessoas.insert(tk.END, f"{k} - {status}{tag_str}")

    def cadastrar_pessoa(self):
        p = self.entry_add_pessoa.get().strip()
        if not p.isdigit(): return messagebox.showerror("Erro", "Nº inválido.")
        pessoas_db = self.current_state.get('pessoas', {})
        if p in pessoas_db: messagebox.showinfo("Aviso", "Pessoa já cadastrada.")
        else:
            pessoas_db[p] = {"ativo": True}
            self.current_state['pessoas'] = pessoas_db
            self.save_state(self.current_state)
            self.atualizar_lista_pessoas()
            self.entry_add_pessoa.delete(0, tk.END)

    def alternar_status_pessoa(self):
        selection = self.listbox_pessoas.curselection()
        if not selection: return
        p = self.listbox_pessoas.get(selection[0]).split(" - ")[0]
        pessoas_db = self.current_state.get('pessoas', {})
        if p in pessoas_db:
            pessoas_db[p]['ativo'] = not pessoas_db[p].get('ativo', True)
            self.save_state(self.current_state)
            self.atualizar_lista_pessoas()

    def alternar_po_pessoa(self):
        selection = self.listbox_pessoas.curselection()
        if not selection: return
        p = self.listbox_pessoas.get(selection[0]).split(" - ")[0]
        pessoas_db = self.current_state.get('pessoas', {})
        if p in pessoas_db:
            pessoas_db[p]['is_po'] = not pessoas_db[p].get('is_po', False)
            self.save_state(self.current_state)
            self.atualizar_lista_pessoas()

    def alternar_sargentiacao_pessoa(self):
        selection = self.listbox_pessoas.curselection()
        if not selection: return
        p = self.listbox_pessoas.get(selection[0]).split(" - ")[0]
        pessoas_db = self.current_state.get('pessoas', {})
        if p in pessoas_db:
            pessoas_db[p]['is_sargentiacao'] = not pessoas_db[p].get('is_sargentiacao', False)
            self.save_state(self.current_state)
            self.atualizar_lista_pessoas()

    def excluir_pessoa(self):
        selection = self.listbox_pessoas.curselection()
        if not selection: return
        p = self.listbox_pessoas.get(selection[0]).split(" - ")[0]
        if messagebox.askyesno("Excluir", f"Tem certeza que deseja apagar o registro de {p}?"):
            pessoas_db = self.current_state.get('pessoas', {})
            if p in pessoas_db:
                del pessoas_db[p]
                self.save_state(self.current_state)
                self.atualizar_lista_pessoas()

    def atualizar_lista_dispensas(self):
        self.listbox_dispensas.delete(0, tk.END)
        for pessoa, dates in self.dispensas.items():
            for inicio, fim in dates:
                self.listbox_dispensas.insert(tk.END, f"P {pessoa}: {inicio.strftime('%d/%m/%Y')} a {fim.strftime('%d/%m/%Y')}")

    def adicionar_dispensa(self):
        p_str = self.entry_disp_pessoa.get().strip()
        inicio = self.parse_dt_local(self.entry_disp_inicio.get().strip())
        fim = self.parse_dt_local(self.entry_disp_fim.get().strip())
        if not p_str.isdigit() or not inicio or not fim or inicio > fim: return
        pessoa = int(p_str)
        if pessoa not in self.dispensas: self.dispensas[pessoa] = []
        self.dispensas[pessoa].append((inicio, fim))
        
        self._sync_dispensas_to_state()
        self.save_state(self.current_state)
        
        self.atualizar_lista_dispensas()
        self.entry_disp_pessoa.delete(0, tk.END)
        self.entry_disp_inicio.delete(0, tk.END)
        self.entry_disp_fim.delete(0, tk.END)

    def remover_dispensa(self):
        selection = self.listbox_dispensas.curselection()
        if not selection: return
        try:
            item_text = self.listbox_dispensas.get(selection[0])
            pessoa = int(item_text.split(":")[0].replace("P ", "").strip())
            # Vamos extrair as datas da string para remover exatamente esta dispensa
            # Formato: "P 301: 10/05/2026 a 20/05/2026"
            datas_str = item_text.split(":")[1].strip()
            inicio_str, fim_str = datas_str.split(" a ")
            inicio = datetime.datetime.strptime(inicio_str, "%d/%m/%Y").date()
            fim = datetime.datetime.strptime(fim_str, "%d/%m/%Y").date()
            
            if pessoa in self.dispensas: 
                if (inicio, fim) in self.dispensas[pessoa]:
                    self.dispensas[pessoa].remove((inicio, fim))
                if len(self.dispensas[pessoa]) == 0:
                    del self.dispensas[pessoa]
                
                self._sync_dispensas_to_state()
                self.save_state(self.current_state)
            
            self.atualizar_lista_dispensas()
        except Exception as e: 
            print("Erro ao remover:", e)

    # --- GERADOR LOGIC ---
    def atualizar_data_alvo(self):
        prox_dia = self.get_next_date()
        dias_nomes = ["Segunda", "Terça", "Quarta", "Quinta", "Sexta", "Sábado", "Domingo"]
        nome_dia = dias_nomes[prox_dia.weekday()]
        self.lbl_target_date.config(text=f"Data Alvo: {prox_dia.strftime('%d/%m/%Y')} ({nome_dia})")
        
        self.var_guarda.set(False)
        self.var_plantao.set(True)
        self.var_apoio.set(False)
        
        if hasattr(self, 'manual_entries'):
            for key, entry in self.manual_entries.items():
                entry.delete(0, tk.END)

    def gerar_escala(self):
        target_date = self.get_next_date()
        has_g = self.var_guarda.get()
        has_p = self.var_plantao.get()
        has_a = self.var_apoio.get()
        
        if not has_g and not has_p and not has_a:
            return messagebox.showerror("Erro", "Selecione pelo menos uma função.")

        try:
            target_str = target_date.strftime("%Y-%m-%d")
            
            guarda_comp_hoje = {}
            for k, entry in self.manual_entries.items():
                guarda_comp_hoje[k] = entry.get().strip() or "-"
            
            # Categorias
            for k, cb in self.manual_categories.items():
                guarda_comp_hoje[k + "_cat"] = cb.get()

            result, new_state = generate_daily_schedule(
                target_date, has_g, has_p, has_a, self.dispensas, self.current_state
            )
            
            result['guarda_comp'] = guarda_comp_hoje
            if new_state.get('historico_escalas'):
                new_state['historico_escalas'][-1]['guarda_comp'] = guarda_comp_hoje
            
            self.schedule_result = result
            self.preview_state = new_state
            
            self.text_details.delete("1.0", tk.END)
            self.text_details.insert(tk.END, f"--- ESCALA PRÉVIA: {target_date.strftime('%d/%m/%Y')} ---\n\n")
            
            self.text_details.insert(tk.END, "--- PREENCHIMENTO MANUAL ---\n")
            for label, key in self.funcoes_list:
                val = guarda_comp_hoje.get(key, "-")
                self.text_details.insert(tk.END, f"{label}: {val}\n")
            self.text_details.insert(tk.END, "\n")
            
            if has_g:
                self.text_details.insert(tk.END, f"GUARDA ({len(result['guarda'])}): {', '.join(result['guarda'])}\n\n")
                
            if has_p: self.text_details.insert(tk.END, f"PLANTÃO ({len(result['plantao'])}): {', '.join(result['plantao'])}\n\n")
            if has_a: self.text_details.insert(tk.END, f"APOIO HT/PARIA ({len(result['apoio'])}): {', '.join(result['apoio'])}\n\n")
            
            self.btn_confirm.config(state=tk.NORMAL)
            self.btn_print.config(state=tk.DISABLED)
            
        except Exception as e:
            messagebox.showerror("Erro", f"Erro:\n{str(e)}")

    def confirmar_escala(self):
        if not self.preview_state: return
        if messagebox.askyesno("Confirmar", "Confirmar esta escala e avançar o dia?"):
            self.current_state = self.preview_state
            
            # Incrementar automaticamente Nr Aditamento e Boletim
            try:
                curr_adit = self.current_state.get('nr_aditamento', '0')
                if curr_adit.isdigit():
                    self.current_state['nr_aditamento'] = str(int(curr_adit) + 1)
                    if hasattr(self, 'entry_nr_adit'):
                        self.entry_nr_adit.delete(0, tk.END)
                        self.entry_nr_adit.insert(0, self.current_state['nr_aditamento'])
                
                curr_bol = self.current_state.get('nr_boletim', '0')
                if curr_bol.isdigit():
                    self.current_state['nr_boletim'] = str(int(curr_bol) + 1)
                    if hasattr(self, 'entry_nr_bol'):
                        self.entry_nr_bol.delete(0, tk.END)
                        self.entry_nr_bol.insert(0, self.current_state['nr_boletim'])
            except Exception as e:
                print("Erro ao incrementar números:", e)

            self.save_state(self.current_state)
            self.btn_confirm.config(state=tk.DISABLED)
            self.btn_print.config(state=tk.NORMAL)
            self.atualizar_data_alvo()
            self.atualizar_lista_historico()
            self.atualizar_ranking()
            if messagebox.askyesno("Imprimir", "Deseja imprimir agora?"):
                self.imprimir_escala(self.schedule_result)

    def imprimir_escala(self, item=None):
        if not item: item = self.schedule_result
        if not item: return
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
            from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn

            def set_cell_background(cell, color_hex):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), color_hex)
                tcPr.append(shd)

            docx_file = "aditamento.docx"
            document = Document()
            
            sections = document.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.8)
                section.right_margin = Inches(0.8)
                
            style = document.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(11)

            d = item['data']
            if isinstance(d, str):
                d_dt = parse_dt(d)
                ds = ["Segunda – feira", "Terça – feira", "Quarta – feira", "Quinta – feira", "Sexta – feira", "Sábado", "Domingo"][d_dt.weekday()]
            else:
                d_dt = d
                ds = ["Segunda – feira", "Terça – feira", "Quarta – feira", "Quinta – feira", "Sexta – feira", "Sábado", "Domingo"][item['dia_semana']]

            meses = {1:"janeiro", 2:"fevereiro", 3:"março", 4:"abril", 5:"maio", 6:"junho", 7:"julho", 8:"agosto", 9:"setembro", 10:"outubro", 11:"novembro", 12:"dezembro"}
            d_str = f"{d_dt.day} de {meses[d_dt.month]} de {d_dt.year}"
            
            gc = item.get('guarda_comp', {})
            
            gda_str = " - ".join(item.get('guarda', [])) if item.get('guarda') else "-"
            plantao_str = " - ".join(item.get('plantao', [])) if item.get('plantao') else "-"
            apoio_str = " - ".join(item.get('apoio', [])) if item.get('apoio') else "-"

            # Dados da Unidade Dinâmica
            unidade_key = self.current_state.get('unidade', 'BC')
            udata = self.UNIDADES_DATA.get(unidade_key, self.UNIDADES_DATA['BC'])
            u_nome = udata['nome']
            u_sigla = udata['sigla']
            u_sigla_doc = udata['sigla_doc']



            # --- BRASÃO ---
            brasao_path = self.get_resource_path("brasao.png")
            if os.path.exists(brasao_path):
                p_img = document.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.space_after = Pt(0)
                run_img = p_img.add_run()
                run_img.add_picture(brasao_path, width=Inches(0.7))

            # --- TABELA PARA TEXTO E VISTO ---
            header_table = document.add_table(rows=1, cols=3)
            header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            header_table.autofit = False
            
            # Definir larguras fixas (Total 6.9")
            # Col 0: 0.7, Col 1: 5.5, Col 2: 0.7
            c_left = header_table.cell(0, 0)
            c_mid = header_table.cell(0, 1)
            c_right = header_table.cell(0, 2)
            c_left.width = Inches(1.2)
            c_mid.width = Inches(4.5)
            c_right.width = Inches(1.2)
            
            # Célula central: Títulos
            p_mid = c_mid.paragraphs[0]
            p_mid.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_mid.paragraph_format.line_spacing = Pt(11)
            p_mid.paragraph_format.space_after = Pt(0)
            
            run_h = p_mid.add_run("MINISTÉRIO DA DEFESA\nEXÉRCITO BRASILEIRO\n17º GRUPO DE ARTILHARIA DE CAMPANHA\n(6º Regimento de Artilharia Montada/1915)\nGRUPO JERÔNIMO DE ALBUQUERQUE")
            run_h.bold = True
            run_h.font.size = Pt(11)

            # Célula da direita: Quadro Visto Sgte
            tc = c_right._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border in ['top', 'left', 'bottom', 'right']:
                element = OxmlElement(f'w:{border}')
                element.set(qn('w:val'), 'single')
                element.set(qn('w:sz'), '4')
                element.set(qn('w:space'), '0')
                element.set(qn('w:color'), '000000')
                tcBorders.append(element)
            tcPr.append(tcBorders)
            
            p_v = c_right.paragraphs[0]
            p_v.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_v.paragraph_format.space_before = Pt(35)
            p_v.paragraph_format.space_after = Pt(2)
            p_v.add_run("__________\n").font.size = Pt(8)
            run_vs = p_v.add_run("Visto Sgte")
            run_vs.bold = True
            run_vs.font.size = Pt(10)
            run_vs.underline = True

            # --- TEXTOS ABAIXO DO CABEÇALHO ---
            p2 = document.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.paragraph_format.space_before = Pt(5)
            p2.paragraph_format.space_after = Pt(0)
            p2.add_run(f"Quartel em Natal /RN, {d_str}.\n({ds})")

            nr_adit = self.current_state.get('nr_aditamento', '___')
            nr_bol = self.current_state.get('nr_boletim', '___')
            current_year = datetime.datetime.now().year
            p4 = document.add_paragraph()
            p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p4.add_run(f"ADITAMENTO AO BOLETIM INTERNO DA {u_nome} Nr {nr_adit}/{current_year}, referente ao BOLETIM INTERNO Nr {nr_bol}/{current_year}, do 17º GAC.").bold = True
            
            p_conhecimento = document.add_paragraph()
            p_conhecimento.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_conhecimento.paragraph_format.space_after = Pt(0)
            p_conhecimento.add_run("Para conhecimento desta Subunidade e devida execução, publico o seguinte:")

            p4 = document.add_paragraph()
            p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p4.paragraph_format.space_after = Pt(0)
            run4 = p4.add_run("1ª Parte:\nSERVIÇOS DIÁRIOS")
            run4.bold = True
            
            document.add_paragraph()

            table = document.add_table(rows=15, cols=3)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = 'Table Grid'
            
            # Ajustar larguras das colunas da tabela de serviço
            table.autofit = False
            widths_main = [Inches(1.5), Inches(0.9), Inches(4.5)]
            for r in table.rows:
                for idx, w in enumerate(widths_main):
                    cell = r.cells[idx]
                    cell.width = w
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            
            # Cabeçalho: Serviço Externo
            row_ext = table.rows[0]
            c_ext = row_ext.cells[0]
            c_ext.merge(row_ext.cells[1]).merge(row_ext.cells[2])
            ds_cap = ds.replace("feira", "Feira")
            c_ext.text = f"Servico Externo para {ds_cap}, {d_str}."
            c_ext.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            c_ext.paragraphs[0].runs[0].bold = True
            set_cell_background(c_ext, "D9D9D9")
            
            data_rows = [
                ("MOT VILA", "SD EP", gc.get('mot_vila', '-')),
                ("GDA VILA", "SD EP", gc.get('gda_vila', '-'))
            ]
            
            for i, (f1, f2, f3) in enumerate(data_rows):
                row = table.rows[i+1]
                row.cells[0].text = f1; set_cell_background(row.cells[0], "D9D9D9")
                row.cells[1].text = f2; set_cell_background(row.cells[1], "D9D9D9")
                row.cells[2].text = f3
            
            # Cabeçalho: Serviço Interno
            row_int = table.rows[3]
            c_int = row_int.cells[0]
            c_int.merge(row_int.cells[1]).merge(row_int.cells[2])
            c_int.text = f"Serviço Interno para {ds_cap}, {d_str}."
            c_int.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            c_int.paragraphs[0].runs[0].bold = True
            set_cell_background(c_int, "D9D9D9")
            
            data_rows_in = [
                ("OF DIA", "1º TEN", gc.get('of_dia', '-')),
                ("ADJ OF DIA", "2º SGT", gc.get('adj_of_dia', '-')),
                (f"SGT DIA {u_sigla}", "3º SGT", gc.get('sgt_dia_bia_c', '-')),
                (f"CB DIA {u_sigla}", "CB EP", gc.get('cb_dia_bia_c', '-')),
                ("MOT DIA", "CB CET", gc.get('mot_dia', '-')),
                ("PADIOLEIRO", "SD EP", gc.get('padioleiro', '-')),
                ("SOMBRA", "SD EP", gc.get('sombra', '-')),
                ("GDA QTEL", "SD EV", gda_str),
                ("PLANTÕES", "SD EV", plantao_str),
                ("APOIO PRAIA/HT", "SD EV", apoio_str)
            ]
            
            for i, (f1, f2, f3) in enumerate(data_rows_in):
                row = table.rows[i+4]
                row.cells[0].text = f1; set_cell_background(row.cells[0], "D9D9D9")
                row.cells[1].text = f2; set_cell_background(row.cells[1], "D9D9D9")
                row.cells[2].text = f3

            # Parada Diária
            row_parada = table.rows[14]
            row_parada.cells[0].text = "PARADA DIÁRIA"; set_cell_background(row_parada.cells[0], "D9D9D9")
            row_parada.cells[1].text = "-"; set_cell_background(row_parada.cells[1], "D9D9D9")
            row_parada.cells[2].text = "9h30min"

            # Ajustar larguras e formatação final de todas as células da tabela principal
            widths_main = [Inches(1.5), Inches(0.9), Inches(4.5)]
            for row in table.rows:
                for idx, w in enumerate(widths_main):
                    cell = row.cells[idx]
                    cell.width = w
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.paragraph_format.space_before = Pt(6)
                        p.paragraph_format.space_after = Pt(6)
                        for run in p.runs:
                            run.font.size = Pt(11)
            document.add_paragraph()
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run("2ª Parte\nINSTRUÇÃO").bold = True
            document.add_paragraph("- Sem Alteração.")
            
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run("3ª Parte\nASSUNTOS GERAIS E ADMINISTRATIVOS").bold = True
            
            document.add_paragraph("1. ASSUNTOS GERAIS").runs[0].bold = True
            document.add_paragraph("- Sem Alteração.")
            
            document.add_paragraph("2. ASSUNTOS ADMINISTRATIVOS").runs[0].bold = True
            document.add_paragraph("- Sem Alteração.")
            
            p = document.add_paragraph(f"- INÍCIO DO EXPEDIENTE PARA {ds.upper()}, {d_str.upper()}")
            p_small = document.add_paragraph("a. Atividade: TFM\n- Início do Expediente para OF/ST/SGT: 07h30min pronto no Campo de Futebol: Unif 14º.\n- Início do Expediente para CB/SD EP: 07h30min pronto no Campo de Futebol: Unif 14º.\n- Início do Expediente para SD EV: 06h45min pronto na SU: Unif 14º.")
            p_small.runs[0].font.size = Pt(9)
            
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run("4ª Parte\nJUSTIÇA E DISCIPLINA").bold = True
            document.add_paragraph("- Sem Alteração.")
            
            document.add_paragraph("\n")
            
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"{self.current_state.get('nome_cmt', 'RENAN LOUREIRO LENTZ - Cap')}\nComandante da {u_nome.title()}")
            run.bold = True
            
            document.add_page_break()

            num_words = {
                0: "ZERO", 1: "UM", 2: "DOIS", 3: "TRÊS", 4: "QUATRO", 5: "CINCO", 
                6: "SEIS", 7: "SETE", 8: "OITO", 9: "NOVE", 10: "DEZ", 
                11: "ONZE", 12: "DOZE", 13: "TREZE", 14: "QUATORZE", 15: "QUINZE", 
                16: "DEZESSEIS", 17: "DEZESSETE", 18: "DEZOITO", 19: "DEZENOVE", 20: "VINTE",
                21: "VINTE E UM", 22: "VINTE E DOIS", 23: "VINTE E TRÊS", 24: "VINTE E QUATRO", 25: "VINTE E CINCO",
                26: "VINTE E SEIS", 27: "VINTE E SETE", 28: "VINTE E OITO", 29: "VINTE E NOVE", 30: "TRINTA",
                31: "TRINTA E UM", 32: "TRINTA E DOIS", 33: "TRINTA E TRÊS", 34: "TRINTA E QUATRO", 35: "TRINTA E CINCO",
                36: "TRINTA E SEIS", 37: "TRINTA E SETE", 38: "TRINTA E OITO", 39: "TRINTA E NOVE", 40: "QUARENTA",
                41: "QUARENTA E UM", 42: "QUARENTA E DOIS", 43: "QUARENTA E TRÊS", 44: "QUARENTA E QUATRO", 45: "QUARENTA E CINCO",
                46: "QUARENTA E SEIS", 47: "QUARENTA E SETE", 48: "QUARENTA E OITO", 49: "QUARENTA E NOVE", 50: "CINQUENTA"
            }
            
            plantoes = item.get('plantao', [])
            guardas = item.get('guarda', [])
            apoios = item.get('apoio', [])
            
            of_dia = gc.get('of_dia', '')
            sgt_dia_bc = gc.get('sgt_dia_bia_c', '')
            cb_dia_bc = gc.get('cb_dia_bia_c', '')
            adj_of_dia = gc.get('adj_of_dia', '')
            mot_dia = gc.get('mot_dia', '')
            padioleiro = gc.get('padioleiro', '')
            sombra = gc.get('sombra', '')
            mot_vila = gc.get('mot_vila', '')
            gda_vila = gc.get('gda_vila', '')
            
            def count_p(val):
                if not val or val == "-": return 0
                return len([x for x in val.replace(' - ', '-').split('-') if x.strip()])

            total_forma = len(plantoes)
            total_forma += count_p(gc.get('of_dia'))
            total_forma += count_p(gc.get('sgt_dia_bia_c'))
            total_forma += count_p(gc.get('cb_dia_bia_c'))
            
            # Outros destinos: todos os campos manuais que não estão na "Forma"
            total_outros = len(guardas) + len(apoios)
            outros_keys = ["adj_of_dia", "mot_dia", "padioleiro", "sombra", "mot_vila", "gda_vila", 
                           "cmt_gda", "cmt_gda_vila", "sgt_dia_bia_o", "cb_gda_qtel", "cb_gda_vila", 
                           "cb_dia_bia_o", "mot_sup_dia", "gda_qtel_ep", "reforco_ep", "permanencia_ht", "reforco_ev"]
            for k in outros_keys:
                total_outros += count_p(gc.get(k))
            
            total_geral = total_forma + total_outros
            
            table_pernoite = document.add_table(rows=18, cols=5)
            table_pernoite.alignment = WD_TABLE_ALIGNMENT.CENTER
            table_pernoite.style = 'Table Grid'
            
            # Ajustar larguras para acomodar a coluna SOMA e Punidos
            # Col 0 (GRAD/PROC): 0.8, Col 1-3 (Texto): 4.7, Col 4 (SOMA/TÉRMINO): 1.0
            widths_p = [Inches(0.8), Inches(1.5), Inches(1.5), Inches(1.5), Inches(1.0)]
            for r_idx in range(18):
                for c_idx in range(5):
                    table_pernoite.cell(r_idx, c_idx).width = widths_p[c_idx]

            def merge_and_set(r, c1, c2, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, bg=None, font_size=9):
                cell = table_pernoite.cell(r, c1)
                if c1 != c2:
                    cell.merge(table_pernoite.cell(r, c2))
                cell.text = text
                p = cell.paragraphs[0]
                p.alignment = align
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                for r_run in p.runs: 
                    r_run.bold = bold
                    r_run.font.size = Pt(font_size)
                if bg:
                    set_cell_background(cell, bg)
                return cell
            
            def count_p(val):
                if not val or val == "-" or val == "": return 0
                return len([x for x in val.replace(' - ', '-').replace(', ', '-').replace('\n', '-').split('-') if x.strip()])
                
            c_of_dia = count_p(gc.get('of_dia'))
            c_sgt_dia = count_p(gc.get('sgt_dia_bia_c'))
            c_cb_dia = count_p(gc.get('cb_dia_bia_c'))
            c_plantoes = len(plantoes)
            total_forma = c_of_dia + c_sgt_dia + c_cb_dia + c_plantoes
            
            c0 = table_pernoite.cell(0, 0)
            c0.text = "Visto:\n\n_________\nCmt SU"
            c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in c0.paragraphs[0].runs: r.font.size = Pt(8)
            
            merge_and_set(0, 1, 4, f"MINISTÉRIO DA DEFESA\nEXÉRCITO BRASILEIRO\n17º GRUPO DE ARTILHARIA DE CAMPANHA\n{u_nome.title()}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            
            d_str_2 = f"{d_dt.day:02d} de {meses[d_dt.month]} de {d_dt.year}"
            merge_and_set(1, 0, 4, f"Controle de Efetivo para {ds}, {d_str_2}.", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg="D9D9D9")
            
            merge_and_set(2, 0, 0, "GRAD", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg="D9D9D9")
            merge_and_set(2, 1, 3, "EM FORMA", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg="D9D9D9")
            merge_and_set(2, 4, 4, "SOMA", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg="D9D9D9")
            
            merge_and_set(3, 0, 0, "TEN", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            merge_and_set(3, 1, 3, f"OF DIA: {gc.get('of_dia','')}", bold=True)
            merge_and_set(3, 4, 4, f"{c_of_dia:02d}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            
            merge_and_set(4, 0, 0, "SGT", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            merge_and_set(4, 1, 3, f"SGT DIA {u_sigla_doc}: {gc.get('sgt_dia_bia_c','')}", bold=True)
            merge_and_set(4, 4, 4, f"{c_sgt_dia:02d}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            
            merge_and_set(5, 0, 0, "SD EP", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            merge_and_set(5, 1, 3, f"CB DIA {u_sigla_doc}: {gc.get('cb_dia_bia_c','')}", bold=True)
            merge_and_set(5, 4, 4, f"{c_cb_dia:02d}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            
            merge_and_set(6, 0, 0, "SD EV", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            merge_and_set(6, 1, 3, f"PLANTÕES: {' - '.join(plantoes)}", bold=True)
            merge_and_set(6, 4, 4, f"{c_plantoes:02d}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            
            merge_and_set(7, 0, 3, f"TOTAL EM FORMA: {num_words.get(total_forma, str(total_forma))}", bold=True)
            merge_and_set(7, 4, 4, f"{total_forma:02d}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

            merge_and_set(8, 0, 4, "PUNIDOS DISCIPLINARMENTE", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg="D9D9D9")
            merge_and_set(9, 0, 0, "PROC.", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg="D9D9D9")
            merge_and_set(9, 1, 1, "GRAD/NOME", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg="D9D9D9")
            merge_and_set(9, 2, 2, "TIPO", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg="D9D9D9")
            merge_and_set(9, 3, 3, "INÍCIO", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg="D9D9D9")
            merge_and_set(9, 4, 4, "TÉRMINO", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg="D9D9D9")
            
            merge_and_set(10, 0, 0, "-", align=WD_ALIGN_PARAGRAPH.CENTER)
            merge_and_set(10, 1, 1, "-", align=WD_ALIGN_PARAGRAPH.CENTER)
            merge_and_set(10, 2, 2, "-", align=WD_ALIGN_PARAGRAPH.CENTER)
            merge_and_set(10, 3, 3, "-", align=WD_ALIGN_PARAGRAPH.CENTER)
            merge_and_set(10, 4, 4, "-", align=WD_ALIGN_PARAGRAPH.CENTER)
            
            merge_and_set(11, 0, 3, "EM OUTROS DESTINOS", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg="D9D9D9")
            merge_and_set(11, 4, 4, "SOMA", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg="D9D9D9")
            
            # Lógica de categorias e contagem para Outros Destinos
            c_sgt_outros = count_p(gc.get('adj_of_dia')) + count_p(gc.get('cmt_gda')) + count_p(gc.get('cmt_gda_vila')) + count_p(gc.get('sgt_dia_bia_o'))
            
            mot_dia = gc.get('mot_dia', '')
            mot_vila = gc.get('mot_vila', '')
            padioleiro = gc.get('padioleiro', '')
            
            mot_dia_cb = mot_dia if gc.get('mot_dia_cat') == 'CB' else ''
            mot_vila_cb = mot_vila if gc.get('mot_vila_cat') == 'CB' else ''
            padioleiro_cb = padioleiro if gc.get('padioleiro_cat') == 'CB' else ''
            c_cb_outros = count_p(gc.get('cb_gda_qtel')) + count_p(gc.get('cb_gda_vila')) + count_p(gc.get('cb_dia_bia_o')) + count_p(mot_dia_cb) + count_p(mot_vila_cb) + count_p(gc.get('mot_sup_dia')) + count_p(padioleiro_cb)

            mot_dia_ep = mot_dia if gc.get('mot_dia_cat') == 'SD EP' else ''
            mot_vila_ep = mot_vila if gc.get('mot_vila_cat') == 'SD EP' else ''
            padioleiro_ep = padioleiro if gc.get('padioleiro_cat') == 'SD EP' else ''
            c_ep_outros = count_p(mot_vila_ep) + count_p(mot_dia_ep) + count_p(gc.get('gda_qtel_ep')) + count_p(gc.get('reforco_ep')) + count_p(gc.get('gda_vila')) + count_p(gc.get('permanencia_ht'))

            mot_vila_ev = mot_vila if gc.get('mot_vila_cat') == 'SD EV' else ''
            padioleiro_ev = padioleiro if gc.get('padioleiro_cat') == 'SD EV' else ''
            c_ev_outros = len(guardas) + count_p(gc.get('gda_vila')) + len(apoios) + count_p(padioleiro_ev) + count_p(gc.get('reforco_ev')) + count_p(mot_vila_ev)
            
            total_outros = c_sgt_outros + c_cb_outros + c_ep_outros + c_ev_outros
            total_geral = total_forma + total_outros

            merge_and_set(12, 0, 0, "SGT", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            merge_and_set(12, 1, 3, f"ADJ OF DIA: {gc.get('adj_of_dia','')}\nCMT GDA: {gc.get('cmt_gda','')}\nCMT GDA VILA: {gc.get('cmt_gda_vila','')}\nSGT DIA 1ª / 2ª BIA O: {gc.get('sgt_dia_bia_o','')}", bold=True)
            merge_and_set(12, 4, 4, f"{c_sgt_outros:02d}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

            merge_and_set(13, 0, 0, "CB EP", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            merge_and_set(13, 1, 3, f"CB GDA QTEL: {gc.get('cb_gda_qtel','')}\nCB GDA VILA: {gc.get('cb_gda_vila','')}\nCB DIA 1ª BIA O: {gc.get('cb_dia_bia_o','')}\nCB DIA 2ª BIA O: \nMOT DIA: {mot_dia_cb}\nMOT VILA: {mot_vila_cb}\nMOT SUP DIA: {gc.get('mot_sup_dia','')}\nPADIOLEIRO: {padioleiro_cb}", bold=True)
            merge_and_set(13, 4, 4, f"{c_cb_outros:02d}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

            merge_and_set(14, 0, 0, "SD EP", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            merge_and_set(14, 1, 3, f"MOT VILA: {mot_vila_ep}\nMOT DIA: {mot_dia_ep}\nGDA QTEL: {gc.get('gda_qtel_ep','')}\nREFORÇO: {gc.get('reforco_ep','')}\nGDA VILA: {gc.get('gda_vila','')}\nPERMANÊNCIA HT: {gc.get('permanencia_ht','')}", bold=True)
            merge_and_set(14, 4, 4, f"{c_ep_outros:02d}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

            merge_and_set(15, 0, 0, "SD EV", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            merge_and_set(15, 1, 3, f"GDA QTEL: {' - '.join(guardas)}\nGDA VILA: {gc.get('gda_vila','')}\nPERMANÊNCIA HT/PRAIA: {' - '.join(apoios)}\nPADIOLEIRO: {padioleiro_ev}\nREFORÇO: {gc.get('reforco_ev','')}\nMOT VILA: {mot_vila_ev}", bold=True)
            merge_and_set(15, 4, 4, f"{c_ev_outros:02d}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            
            merge_and_set(16, 0, 3, f"TOTAL EM OUTROS DESTINOS: {num_words.get(total_outros, str(total_outros))}", bold=True)
            merge_and_set(16, 4, 4, f"{total_outros:02d}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            
            merge_and_set(17, 0, 3, f"TOTAL GERAL: {num_words.get(total_geral, str(total_geral))}", bold=True)
            merge_and_set(17, 4, 4, f"{total_geral:02d}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            
            # Ajustar larguras da tabela pernoite para totalizar 6.9"
            widths_p = [Inches(0.8), Inches(1.7), Inches(1.7), Inches(1.7), Inches(1.0)]
            for row in table_pernoite.rows:
                for idx, width in enumerate(widths_p):
                    try: 
                        cell = row.cells[idx]
                        cell.width = width
                    except: pass

            # --- TABELA DE ASSINATURAS E ALTERAÇÃO ---
            table_sig = document.add_table(rows=2, cols=3)
            table_sig.alignment = WD_TABLE_ALIGNMENT.CENTER
            table_sig.style = 'Table Grid'
            table_sig.autofit = False
            
            # Ajustar larguras para totalizar 6.9" e alinhar com a de cima
            widths_sig = [Inches(1.5), Inches(3.9), Inches(1.5)]
            for r_idx in range(2):
                for c_idx in range(3):
                    table_sig.cell(r_idx, c_idx).width = widths_sig[c_idx]

            # Visto Sgt Dia
            c_sgt = table_sig.cell(0, 0)
            p_sgt = c_sgt.paragraphs[0]
            p_sgt.text = "Visto:"
            p_sgt.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_sgt.runs[0].font.size = Pt(8)
            
            p_sgt_line = c_sgt.add_paragraph()
            p_sgt_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_sgt_line.paragraph_format.space_before = Pt(25)
            p_sgt_line.add_run("_________").font.size = Pt(8)
            p_sgt_line.add_run("\nSgt Dia").font.size = Pt(9)
            p_sgt_line.runs[1].bold = True
            p_sgt_line.runs[1].underline = True
            
            # Centro (Quartel + Sargenteante)
            c_mid = table_sig.cell(0, 1)
            nome_sgte = self.current_state.get('nome_sgte', 'HEBERT CARLOS VIANA - 2° Sgt')
            
            p_date = c_mid.paragraphs[0]
            p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_date = p_date.add_run(f"Quartel em Natal/RN, {d_str}.")
            r_date.bold = True
            r_date.font.size = Pt(10)
            
            p_name = c_mid.add_paragraph()
            p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_name.paragraph_format.space_before = Pt(10)
            r_name = p_name.add_run(nome_sgte)
            r_name.bold = True
            r_name.font.size = Pt(11)
            
            p_title = c_mid.add_paragraph(f"Sargenteante da {u_nome.title()}")
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_title.paragraph_format.space_after = Pt(0)
            for r in p_title.runs: 
                r.font.size = Pt(9)
                r.underline = True
            
            # Visto Of Dia
            c_of = table_sig.cell(0, 2)
            p_of = c_of.paragraphs[0]
            p_of.text = "Visto:"
            p_of.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_of.runs[0].font.size = Pt(8)
            
            p_of_line = c_of.add_paragraph()
            p_of_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_of_line.paragraph_format.space_before = Pt(25)
            p_of_line.add_run("_________").font.size = Pt(8)
            p_of_line.add_run("\nOf Dia").font.size = Pt(9)
            p_of_line.runs[1].bold = True
            p_of_line.runs[1].underline = True

            # --- SEÇÃO DE ALTERAÇÃO FINAL (Linha 2 da tabela) ---
            c_alt2 = table_sig.cell(1, 0)
            c_alt2.merge(table_sig.cell(1, 1)).merge(table_sig.cell(1, 2))
            
            p_alt2 = c_alt2.paragraphs[0]
            p_alt2.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run_alt2 = p_alt2.add_run("Alteração: Com alteração (      ) Sem alteração (      )")
            run_alt2.bold = True
            run_alt2.font.size = Pt(9)
            p_alt2.paragraph_format.line_spacing = Pt(12)
            
            for _ in range(4):
                # Usar Tab Stop com preenchimento de linha para garantir largura total perfeita
                p_l = c_alt2.add_paragraph()
                p_l.paragraph_format.line_spacing = Pt(12)
                p_l.paragraph_format.space_before = Pt(0)
                p_l.paragraph_format.space_after = Pt(0)
                p_l.paragraph_format.tab_stops.add_tab_stop(Inches(6.6), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.LINES)
                p_l.add_run("\t")
                for r in p_l.runs: r.font.size = Pt(9)
            
            try:
                document.save(docx_file)
                os.startfile(docx_file)
            except Exception as e:
                messagebox.showerror("Erro de Arquivo", f"Não foi possível salvar o documento.\nCertifique-se de que o Word não está aberto com um arquivo de mesmo nome.\n\nErro: {e}")
            
        except ImportError:
            messagebox.showerror("Erro", "A biblioteca python-docx não está instalada.")
        except Exception as e:
            messagebox.showerror("Erro", f"Erro de Impressão:\n{e}")

    # --- HISTORICO LOGIC ---
    def atualizar_lista_historico(self):
        self.listbox_historico.delete(0, tk.END)
        historico = self.current_state.get('historico_escalas', [])
        for i, reg in enumerate(historico):
            dt = parse_dt(reg['data'])
            dt_str = dt.strftime("%d/%m/%Y")
            self.listbox_historico.insert(tk.END, f"{i} - Escala do dia {dt_str}")

    def on_historico_select(self, event):
        selection = self.listbox_historico.curselection()
        if not selection: return
        idx = int(self.listbox_historico.get(selection[0]).split(" - ")[0])
        historico = self.current_state.get('historico_escalas', [])
        if idx >= len(historico): return
        
        item = historico[idx]
        self.text_hist_details.delete("1.0", tk.END)
        dt = parse_dt(item['data'])
        self.text_hist_details.insert(tk.END, f"--- HISTÓRICO: {dt.strftime('%d/%m/%Y')} ---\n\n")
        
        if 'guarda_comp' in item:
            gc = item['guarda_comp']
            self.text_hist_details.insert(tk.END, "--- PREENCHIMENTO MANUAL ---\n")
            for label, key in self.funcoes_list:
                val = gc.get(key, "-")
                self.text_hist_details.insert(tk.END, f"{label}: {val}\n")
            self.text_hist_details.insert(tk.END, "\n")
            
        if item.get('has_guarda'): 
            self.text_hist_details.insert(tk.END, f"GUARDA ({len(item.get('guarda', []))}): {', '.join(item.get('guarda', []))}\n\n")
            
        if item.get('has_plantao'): self.text_hist_details.insert(tk.END, f"PLANTÃO ({len(item['plantao'])}): {', '.join(item['plantao'])}\n\n")
        if item.get('has_apoio'): self.text_hist_details.insert(tk.END, f"APOIO HT/PARIA ({len(item['apoio'])}): {', '.join(item['apoio'])}\n\n")

    def excluir_historico(self):
        selection = self.listbox_historico.curselection()
        if not selection: return
        idx = int(self.listbox_historico.get(selection[0]).split(" - ")[0])
        historico = self.current_state.get('historico_escalas', [])
        
        if idx >= len(historico): return
        
        dt_str = parse_dt(historico[idx]['data']).strftime("%d/%m/%Y")
        if messagebox.askyesno("Excluir", f"Você tem certeza que deseja EXCLUIR a escala do dia {dt_str}?\n\nOs registros de trabalho de todo o efetivo neste dia serão apagados!"):
            del historico[idx]
            self.current_state['historico_escalas'] = historico
            self.save_state(self.current_state)
            self.atualizar_lista_historico()
            self.atualizar_data_alvo()
            self.text_hist_details.delete("1.0", tk.END)
            messagebox.showinfo("Sucesso", "Escala excluída com sucesso! Os serviços das pessoas voltaram como se este dia nunca tivesse existido.")

    def imprimir_historico(self):
        selection = self.listbox_historico.curselection()
        if not selection: 
            return messagebox.showwarning("Aviso", "Selecione uma escala no histórico primeiro.")
        
        idx = int(self.listbox_historico.get(selection[0]).split(" - ")[0])
        historico = self.current_state.get('historico_escalas', [])
        
        if idx >= len(historico): return
        
        self.imprimir_escala(historico[idx])

if __name__ == "__main__":
    root = tk.Tk()
    app = App(root)
    sv_ttk.set_theme("dark")
    root.mainloop()
