import sys
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def create_pernoite():
    document = Document()
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)

    table = document.add_table(rows=20, cols=4)
    table.style = 'Table Grid'

    def merge_and_set(r, c1, c2, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, bg=None):
        cell = table.cell(r, c1)
        if c1 != c2:
            cell.merge(table.cell(r, c2))
        cell.text = text
        p = cell.paragraphs[0]
        p.alignment = align
        if bold:
            for run in p.runs: run.bold = True
        if bg:
            set_cell_background(cell, bg)
        return cell

    c0 = table.cell(0, 0)
    c0.text = "Visto:\n\n\n\nCmt SU"
    c0.paragraphs[0].runs[0].font.size = Pt(8)
    
    merge_and_set(0, 1, 3, "MINISTÉRIO DA DEFESA\nEXÉRCITO BRASILEIRO\n17º GRUPO DE ARTILHARIA DE CAMPANHA\nBateria de Comando", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    merge_and_set(1, 0, 3, "Controle de Efetivo para Quarta-feira, 08 de abril de 2026.", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg="D9D9D9")

    merge_and_set(2, 0, 0, "GRAD", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg="D9D9D9")
    merge_and_set(2, 1, 3, "EM FORMA", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg="D9D9D9")

    merge_and_set(3, 0, 0, "TEN", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    merge_and_set(3, 1, 3, "OF DIA: ALGUEM", bold=True)
    
    merge_and_set(4, 0, 0, "SGT", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    merge_and_set(4, 1, 3, "SGT DIA BC: PESSOA", bold=True)
    
    merge_and_set(5, 0, 0, "SD EP", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    merge_and_set(5, 1, 3, "CB DIA BC: SILVA", bold=True)
    
    merge_and_set(6, 0, 0, "SD EV", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    merge_and_set(6, 1, 3, "PLANTÕES: 301 - 302 - 303", bold=True)

    merge_and_set(7, 0, 3, "TOTAL EM FORMA: OITO", bold=True)

    merge_and_set(8, 0, 3, "PUNIDOS DISCIPLINARMENTE", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg="D9D9D9")

    merge_and_set(9, 0, 0, "PROC.", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg="D9D9D9")
    merge_and_set(9, 1, 1, "GRAD/NOME", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg="D9D9D9")
    merge_and_set(9, 2, 2, "TIPO", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg="D9D9D9")
    merge_and_set(9, 3, 3, "INÍCIO", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg="D9D9D9")

    merge_and_set(10, 0, 0, "-", align=WD_ALIGN_PARAGRAPH.CENTER)
    merge_and_set(10, 1, 1, "-", align=WD_ALIGN_PARAGRAPH.CENTER)
    merge_and_set(10, 2, 2, "-", align=WD_ALIGN_PARAGRAPH.CENTER)
    merge_and_set(10, 3, 3, "-", align=WD_ALIGN_PARAGRAPH.CENTER)

    merge_and_set(11, 0, 3, "EM OUTROS DESTINOS", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg="D9D9D9")

    merge_and_set(12, 0, 0, "SGT", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    merge_and_set(12, 1, 3, "ADJ OF DIA:\nCMT GDA: SÁVIO\nCMT GDA VILA:\nSGT DIA 1\" / 2\" BIA O:", bold=True)

    merge_and_set(13, 0, 0, "CB EP", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    merge_and_set(13, 1, 3, "CB GDA QTEL:\nCB GDA VILA:\nCB DIA 1\" BIA O:\nCB DIA 2\" BIA O:\nMOT DIA:\nMOT VILA:\nMOT SUP DIA:\nPADIOLEIRO: ROQUE", bold=True)

    merge_and_set(14, 0, 0, "SD EP", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    merge_and_set(14, 1, 3, "MOT VILA:\nMOT DIA:\nGDA QTEL:\nREFORÇO:\nGDA VILA:\nPERMANÊNCIA HT:", bold=True)

    merge_and_set(15, 0, 0, "SD EV", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    merge_and_set(15, 1, 3, "GDA QTEL:\nGDA VILA:\nPERMANÊNCIA HT/PRAIA:\nPADIOLEIRO:\nREFORÇO:\nMOT VILA:", bold=True)

    merge_and_set(16, 0, 3, "TOTAL EM OUTROS DESTINOS: ONZE", bold=True)
    merge_and_set(17, 0, 3, "TOTAL GERAL: DEZENOVE", bold=True)

    c18_0 = merge_and_set(18, 0, 0, "Visto:\n\n\nSgt Dia")
    c18_0.paragraphs[0].runs[0].font.size = Pt(8)
    
    merge_and_set(18, 1, 2, "Quartel em Natal/RN, 08 de abril de 2026.\n\nHEBERT CARLOS VIANA - 2° Sgt\nSargenteante da Bateria de Comando", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    c18_3 = merge_and_set(18, 3, 3, "Visto:\n\n\nOf Dia")
    c18_3.paragraphs[0].runs[0].font.size = Pt(8)

    merge_and_set(19, 0, 3, "Alteração: Com alteração (  ) Sem alteração (  )\n\n\n\n")

    widths = [Inches(1.0), Inches(3.5), Inches(1.5), Inches(1.5)]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    document.save("test_pernoite.docx")
    print("Salvo como test_pernoite.docx")

create_pernoite()
