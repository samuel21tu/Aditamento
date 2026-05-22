import os
import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from scheduler import parse_dt
from constants import NUM_WORDS, UNIDADES_DATA

class DocumentGenerator:
    def __init__(self, current_state, get_resource_path_callback):
        self.current_state = current_state
        self.get_resource_path = get_resource_path_callback

    def set_cell_background(self, cell, color_hex):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color_hex)
        tcPr.append(shd)

    def merge_and_set(self, table, r, c1, c2, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, bg=None, font_size=9):
        cell = table.cell(r, c1)
        if c1 != c2:
            cell.merge(table.cell(r, c2))
        cell.text = text
        p = cell.paragraphs[0]
        p.alignment = align
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        for r_run in p.runs: 
            r_run.bold = bold
            r_run.font.size = Pt(font_size)
        if bg:
            self.set_cell_background(cell, bg)
        return cell

    def count_p(self, val):
        if not val or val == "-" or val == "": return 0
        return len([x for x in val.replace(' - ', '-').replace(', ', '-').replace('\n', '-').split('-') if x.strip()])

    def generate(self, item, output_path="aditamento.docx"):
        document = Document()
        
        sections = document.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)
            
        style = document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(11)

        d = item['data']
        if isinstance(d, str):
            d_dt = parse_dt(d)
            ds = ["Segunda – feira", "Terça – feira", "Quarta – feira", "Quinta – feira", "Sexta – feira", "Sábado", "Domingo"][d_dt.weekday()]
        else:
            d_dt = d
            ds = ["Segunda – feira", "Terça – feira", "Quarta – feira", "Quinta – feira", "Sexta – feira", "Sábado", "Domingo"][item['dia_semana']]

        meses = {1:"janeiro", 2:"fevereiro", 3:"março", 4:"abril", 5:"maio", 6:"junho", 7:"julho", 8:"agosto", 9:"setembro", 10:"outubro", 11:"novembro", 12:"dezembro"}
        d_str = f"{d_dt.day} de {meses[d_dt.month]} de {d_dt.year}"
        
        is_vermelha = item.get('sem_expediente', d_dt.weekday() >= 5)
        parada_hora = "07h30min" if is_vermelha else "09h30min"
        tfm_text = "Início do Expediente" if is_vermelha else "TFM"
        local_pronto = "pronto na sua subunidade" if is_vermelha else "pronto no Campo de Futebol"
        uniforme = "Uniforme da seção" if is_vermelha else "Unif 14º"
        
        gc = item.get('guarda_comp', {})
        
        gda_str = " - ".join(item.get('guarda', [])) if item.get('guarda') else "-"
        plantao_str = " - ".join(item.get('plantao', [])) if item.get('plantao') else "-"
        apoio_str = " - ".join(item.get('apoio', [])) if item.get('apoio') else "-"
        sobre_aviso_str = " - ".join(item.get('sobre_aviso', [])) if item.get('sobre_aviso') else "-"

        unidade_key = self.current_state.get('unidade', 'BC')
        udata = UNIDADES_DATA.get(unidade_key, UNIDADES_DATA['BC'])
        u_nome = udata['nome']
        u_sigla = udata['sigla']
        u_sigla_doc = udata['sigla_doc']

        # --- BRASÃO ---
        brasao_path = self.get_resource_path("brasao.png")
        if os.path.exists(brasao_path):
            p_img = document.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_after = Pt(0)
            run_img = p_img.add_run()
            run_img.add_picture(brasao_path, width=Inches(0.7))

        # --- CABEÇALHO ---
        header_table = document.add_table(rows=1, cols=3)
        header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        header_table.autofit = False
        
        c_mid = header_table.cell(0, 1)
        c_right = header_table.cell(0, 2)
        header_table.cell(0, 0).width = Inches(1.2)
        c_mid.width = Inches(4.5)
        c_right.width = Inches(1.2)
        
        p_mid = c_mid.paragraphs[0]
        p_mid.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_mid.paragraph_format.line_spacing = Pt(11)
        p_mid.paragraph_format.space_after = Pt(0)
        
        run_h = p_mid.add_run("MINISTÉRIO DA DEFESA\nEXÉRCITO BRASILEIRO\n17º GRUPO DE ARTILHARIA DE CAMPANHA\n(6º Regimento de Artilharia Montada/1915)\nGRUPO JERÔNIMO DE ALBUQUERQUE")
        run_h.bold = True
        run_h.font.size = Pt(11)

        tc = c_right._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for border in ['top', 'left', 'bottom', 'right']:
            element = OxmlElement(f'w:{border}')
            element.set(qn('w:val'), 'single')
            element.set(qn('w:sz'), '4')
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), '000000')
            tcBorders.append(element)
        tcPr.append(tcBorders)
        
        p_v = c_right.paragraphs[0]
        p_v.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_v.paragraph_format.space_before = Pt(35)
        p_v.paragraph_format.space_after = Pt(2)
        p_v.add_run("__________\n").font.size = Pt(8)
        run_vs = p_v.add_run("Visto Sgte")
        run_vs.bold = True
        run_vs.font.size = Pt(10)
        run_vs.underline = True

        p2 = document.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(5)
        p2.paragraph_format.space_after = Pt(0)
        p2.add_run(f"Quartel em Natal /RN, {d_str}.\n({ds})")

        nr_adit = self.current_state.get('nr_aditamento', '___')
        nr_bol = self.current_state.get('nr_boletim', '___')
        current_year = datetime.datetime.now().year
        p4 = document.add_paragraph()
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p4.add_run(f"ADITAMENTO AO BOLETIM INTERNO DA {u_nome} Nr {nr_adit}/{current_year}, referente ao BOLETIM INTERNO Nr {nr_bol}/{current_year}, do 17º GAC.").bold = True
        
        p_conhecimento = document.add_paragraph()
        p_conhecimento.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_conhecimento.paragraph_format.space_after = Pt(0)
        p_conhecimento.add_run("Para conhecimento desta Subunidade e devida execução, publico o seguinte:")

        p4 = document.add_paragraph()
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p4.paragraph_format.space_after = Pt(0)
        run4 = p4.add_run("1ª Parte:\nSERVIÇOS DIÁRIOS")
        run4.bold = True
        
        document.add_paragraph()

        # --- TABELA DE SERVIÇOS ---
        table = document.add_table(rows=0, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        table.autofit = False
        widths_main = [Inches(1.5), Inches(0.9), Inches(4.5)]
        
        ds_cap = ds.replace("feira", "Feira")
        
        # Filtro de serviços externos
        external_filled = []
        raw_external_rows = [
            ("MOT VILA", "SD EP", gc.get('mot_vila', '-')),
            ("GDA VILA", "SD EP", gc.get('gda_vila', '-'))
        ]
        for f1, f2, f3 in raw_external_rows:
            if f3 and f3.strip() not in ("", "-"):
                external_filled.append((f1, f2, f3))
                
        if external_filled:
            row_ext = table.add_row()
            c_ext = row_ext.cells[0]
            c_ext.merge(row_ext.cells[1]).merge(row_ext.cells[2])
            c_ext.text = f"Servico Externo para {ds_cap}, {d_str}."
            c_ext.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            c_ext.paragraphs[0].runs[0].bold = True
            self.set_cell_background(c_ext, "D9D9D9")
            
            for f1, f2, f3 in external_filled:
                row = table.add_row()
                row.cells[0].text = f1; self.set_cell_background(row.cells[0], "D9D9D9")
                row.cells[1].text = f2; self.set_cell_background(row.cells[1], "D9D9D9")
                row.cells[2].text = f3
                
        # Filtro de serviços internos
        internal_filled = []
        raw_internal_rows = [
            ("OF DIA", "1º TEN", gc.get('of_dia', '-')),
            ("ADJ OF DIA", "2º SGT", gc.get('adj_of_dia', '-')),
            (f"SGT DIA {u_sigla}", "3º SGT", gc.get('sgt_dia_bia_c', '-')),
            (f"CB DIA {u_sigla}", "CB EP", gc.get('cb_dia_bia_c', '-')),
            ("MOT DIA", "CB CET", gc.get('mot_dia', '-')),
            ("PADIOLEIRO", "SD EP", gc.get('padioleiro', '-')),
            ("SOMBRA", "SD EP", gc.get('sombra', '-')),
            ("GDA QTEL", "SD EV", gda_str),
            ("PLANTÕES", "SD EV", plantao_str),
            ("APOIO PRAIA/HT", "SD EV", apoio_str),
            ("SOBRE AVISO", "SD EV", sobre_aviso_str)
        ]
        for f1, f2, f3 in raw_internal_rows:
            if f3 and f3.strip() not in ("", "-"):
                internal_filled.append((f1, f2, f3))
                
        if internal_filled:
            row_int = table.add_row()
            c_int = row_int.cells[0]
            c_int.merge(row_int.cells[1]).merge(row_int.cells[2])
            c_int.text = f"Serviço Interno para {ds_cap}, {d_str}."
            c_int.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            c_int.paragraphs[0].runs[0].bold = True
            self.set_cell_background(c_int, "D9D9D9")
            
            for f1, f2, f3 in internal_filled:
                row = table.add_row()
                row.cells[0].text = f1; self.set_cell_background(row.cells[0], "D9D9D9")
                row.cells[1].text = f2; self.set_cell_background(row.cells[1], "D9D9D9")
                row.cells[2].text = f3

        # Parada Diária
        row_parada = table.add_row()
        row_parada.cells[0].text = "PARADA DIÁRIA"; self.set_cell_background(row_parada.cells[0], "D9D9D9")
        row_parada.cells[1].text = "-"; self.set_cell_background(row_parada.cells[1], "D9D9D9")
        row_parada.cells[2].text = parada_hora

        for row in table.rows:
            for idx, w in enumerate(widths_main):
                cell = row.cells[idx]
                cell.width = w
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                    for run in p.runs:
                        run.font.size = Pt(11)
        
        document.add_paragraph()
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("2ª Parte\nINSTRUÇÃO").bold = True
        document.add_paragraph("- Sem Alteração.")
        
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("3ª Parte\nASSUNTOS GERAIS E ADMINISTRATIVOS").bold = True
        
        document.add_paragraph("1. ASSUNTOS GERAIS").runs[0].bold = True
        document.add_paragraph("- Sem Alteração.")
        
        document.add_paragraph("2. ASSUNTOS ADMINISTRATIVOS").runs[0].bold = True
        document.add_paragraph("- Sem Alteração.")
        
        p = document.add_paragraph(f"- INÍCIO DO EXPEDIENTE PARA {ds.upper()}, {d_str.upper()}")
        p_small = document.add_paragraph(f"a. Atividade: {tfm_text}\n- Início do Expediente para OF/ST/SGT: 07h30min {local_pronto}: {uniforme}.\n- Início do Expediente para CB/SD EP: 07h30min {local_pronto}: {uniforme}.\n- Início do Expediente para SD EV: 06h45min pronto na SU: {uniforme}.")
        p_small.runs[0].font.size = Pt(9)
        
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("4ª Parte\nJUSTIÇA E DISCIPLINA").bold = True
        document.add_paragraph("- Sem Alteração.\n")
        
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{self.current_state.get('nome_cmt', 'RENAN LOUREIRO LENTZ - Cap')}\nComandante da {u_nome.title()}")
        run.bold = True
        
        document.add_page_break()

        # --- PAGO PERNOITE (Controle de Efetivo) ---
        plantoes = item.get('plantao', [])
        guardas = item.get('guarda', [])
        apoios = item.get('apoio', [])
        sobre_avisos = item.get('sobre_aviso', [])

        # Para o Pernoite (Controle de Efetivo / Arranchamento), a data deve ser hoje (ontem em relação à data alvo do aditamento)
        d_dt_yesterday = d_dt - datetime.timedelta(days=1)
        ds_yesterday = ["Segunda – feira", "Terça – feira", "Quarta – feira", "Quinta – feira", "Sexta – feira", "Sábado", "Domingo"][d_dt_yesterday.weekday()]
        d_str_yesterday = f"{d_dt_yesterday.day} de {meses[d_dt_yesterday.month]} de {d_dt_yesterday.year}"
        d_str_2_yesterday = f"{d_dt_yesterday.day:02d} de {meses[d_dt_yesterday.month]} de {d_dt_yesterday.year}"

        table_pernoite = document.add_table(rows=18, cols=5)
        table_pernoite.alignment = WD_TABLE_ALIGNMENT.CENTER
        table_pernoite.style = 'Table Grid'
        
        c_of_dia = self.count_p(gc.get('of_dia'))
        c_sgt_dia = self.count_p(gc.get('sgt_dia_bia_c'))
        c_cb_dia = self.count_p(gc.get('cb_dia_bia_c'))
        c_plantoes = len(plantoes)
        total_forma = c_of_dia + c_sgt_dia + c_cb_dia + c_plantoes
        
        c0 = table_pernoite.cell(0, 0)
        c0.text = "Visto:\n\n_________\nCmt SU"
        c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in c0.paragraphs[0].runs: r.font.size = Pt(8)
        
        self.merge_and_set(table_pernoite, 0, 1, 4, f"MINISTÉRIO DA DEFESA\nEXÉRCITO BRASILEIRO\n17º GRUPO DE ARTILHARIA DE CAMPANHA\n{u_nome.title()}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        
        self.merge_and_set(table_pernoite, 1, 0, 4, f"Controle de Efetivo para {ds_yesterday}, {d_str_2_yesterday}.", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg="D9D9D9")
        
        self.merge_and_set(table_pernoite, 2, 0, 0, "GRAD", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg="D9D9D9")
        self.merge_and_set(table_pernoite, 2, 1, 3, "EM FORMA", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg="D9D9D9")
        self.merge_and_set(table_pernoite, 2, 4, 4, "SOMA", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg="D9D9D9")
        
        self.merge_and_set(table_pernoite, 3, 0, 0, "TEN", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        self.merge_and_set(table_pernoite, 3, 1, 3, f"OF DIA: {gc.get('of_dia','')}", bold=True)
        self.merge_and_set(table_pernoite, 3, 4, 4, f"{c_of_dia:02d}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        
        self.merge_and_set(table_pernoite, 4, 0, 0, "SGT", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        self.merge_and_set(table_pernoite, 4, 1, 3, f"SGT DIA {u_sigla_doc}: {gc.get('sgt_dia_bia_c','')}", bold=True)
        self.merge_and_set(table_pernoite, 4, 4, 4, f"{c_sgt_dia:02d}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        
        self.merge_and_set(table_pernoite, 5, 0, 0, "SD EP", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        self.merge_and_set(table_pernoite, 5, 1, 3, f"CB DIA {u_sigla_doc}: {gc.get('cb_dia_bia_c','')}", bold=True)
        self.merge_and_set(table_pernoite, 5, 4, 4, f"{c_cb_dia:02d}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        
        self.merge_and_set(table_pernoite, 6, 0, 0, "SD EV", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        self.merge_and_set(table_pernoite, 6, 1, 3, f"PLANTÕES: {' - '.join(plantoes)}", bold=True)
        self.merge_and_set(table_pernoite, 6, 4, 4, f"{c_plantoes:02d}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        
        self.merge_and_set(table_pernoite, 7, 0, 3, f"TOTAL EM FORMA: {NUM_WORDS.get(total_forma, str(total_forma))}", bold=True)
        self.merge_and_set(table_pernoite, 7, 4, 4, f"{total_forma:02d}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

        self.merge_and_set(table_pernoite, 8, 0, 4, "PUNIDOS DISCIPLINARMENTE", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg="D9D9D9")
        self.merge_and_set(table_pernoite, 9, 0, 0, "PROC.", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg="D9D9D9")
        self.merge_and_set(table_pernoite, 9, 1, 1, "GRAD/NOME", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg="D9D9D9")
        self.merge_and_set(table_pernoite, 9, 2, 2, "TIPO", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg="D9D9D9")
        self.merge_and_set(table_pernoite, 9, 3, 3, "INÍCIO", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg="D9D9D9")
        self.merge_and_set(table_pernoite, 9, 4, 4, "TÉRMINO", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg="D9D9D9")
        
        for i in range(10, 11):
            for j in range(5): self.merge_and_set(table_pernoite, i, j, j, "-", align=WD_ALIGN_PARAGRAPH.CENTER)
        
        self.merge_and_set(table_pernoite, 11, 0, 3, "EM OUTROS DESTINOS", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg="D9D9D9")
        self.merge_and_set(table_pernoite, 11, 4, 4, "SOMA", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg="D9D9D9")
        
        c_sgt_outros = self.count_p(gc.get('adj_of_dia')) + self.count_p(gc.get('cmt_gda')) + self.count_p(gc.get('cmt_gda_vila')) + self.count_p(gc.get('sgt_dia_bia_o'))
        
        mot_dia_cb = gc.get('mot_dia', '') if gc.get('mot_dia_cat') == 'CB' else ''
        mot_vila_cb = gc.get('mot_vila', '') if gc.get('mot_vila_cat') == 'CB' else ''
        padioleiro_cb = gc.get('padioleiro', '') if gc.get('padioleiro_cat') == 'CB' else ''
        c_cb_outros = self.count_p(gc.get('cb_gda_qtel')) + self.count_p(gc.get('cb_gda_vila')) + self.count_p(gc.get('cb_dia_bia_o')) + self.count_p(mot_dia_cb) + self.count_p(mot_vila_cb) + self.count_p(gc.get('mot_sup_dia')) + self.count_p(padioleiro_cb)

        mot_dia_ep = gc.get('mot_dia', '') if gc.get('mot_dia_cat') == 'SD EP' else ''
        mot_vila_ep = gc.get('mot_vila', '') if gc.get('mot_vila_cat') == 'SD EP' else ''
        padioleiro_ep = gc.get('padioleiro', '') if gc.get('padioleiro_cat') == 'SD EP' else ''
        c_ep_outros = self.count_p(mot_vila_ep) + self.count_p(mot_dia_ep) + self.count_p(gc.get('gda_qtel_ep')) + self.count_p(gc.get('reforco_ep')) + self.count_p(gc.get('gda_vila')) + self.count_p(gc.get('permanencia_ht'))

        mot_vila_ev = gc.get('mot_vila', '') if gc.get('mot_vila_cat') == 'SD EV' else ''
        padioleiro_ev = gc.get('padioleiro', '') if gc.get('padioleiro_cat') == 'SD EV' else ''
        c_ev_outros = len(guardas) + self.count_p(gc.get('gda_vila')) + len(apoios) + self.count_p(padioleiro_ev) + self.count_p(gc.get('reforco_ev')) + self.count_p(mot_vila_ev) + len(sobre_avisos)
        
        total_outros = c_sgt_outros + c_cb_outros + c_ep_outros + c_ev_outros
        total_geral = total_forma + total_outros

        self.merge_and_set(table_pernoite, 12, 0, 0, "SGT", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        self.merge_and_set(table_pernoite, 12, 1, 3, f"ADJ OF DIA: {gc.get('adj_of_dia','')}\nCMT GDA: {gc.get('cmt_gda','')}\nCMT GDA VILA: {gc.get('cmt_gda_vila','')}", bold=True)
        self.merge_and_set(table_pernoite, 12, 4, 4, f"{c_sgt_outros:02d}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

        self.merge_and_set(table_pernoite, 13, 0, 0, "CB EP", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        self.merge_and_set(table_pernoite, 13, 1, 3, f"CB GDA QTEL: {gc.get('cb_gda_qtel','')}\nCB GDA VILA: {gc.get('cb_gda_vila','')}\nMOT DIA: {mot_dia_cb}\nMOT VILA: {mot_vila_cb}\nMOT SUP DIA: {gc.get('mot_sup_dia','')}\nPADIOLEIRO: {padioleiro_cb}", bold=True)
        self.merge_and_set(table_pernoite, 13, 4, 4, f"{c_cb_outros:02d}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

        self.merge_and_set(table_pernoite, 14, 0, 0, "SD EP", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        self.merge_and_set(table_pernoite, 14, 1, 3, f"MOT VILA: {mot_vila_ep}\nMOT DIA: {mot_dia_ep}\nGDA QTEL: {gc.get('gda_qtel_ep','')}\nREFORÇO: {gc.get('reforco_ep','')}\nGDA VILA: {gc.get('gda_vila','')}\nPERMANÊNCIA HT: {gc.get('permanencia_ht','')}", bold=True)
        self.merge_and_set(table_pernoite, 14, 4, 4, f"{c_ep_outros:02d}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

        self.merge_and_set(table_pernoite, 15, 0, 0, "SD EV", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        self.merge_and_set(table_pernoite, 15, 1, 3, f"GDA QTEL: {' - '.join(guardas)}\nGDA VILA: {gc.get('gda_vila','')}\nPERMANÊNCIA HT/PRAIA: {' - '.join(apoios)}\nSOBRE AVISO: {' - '.join(sobre_avisos)}\nPADIOLEIRO: {padioleiro_ev}\nREFORÇO: {gc.get('reforco_ev','')}\nMOT VILA: {mot_vila_ev}", bold=True)
        self.merge_and_set(table_pernoite, 15, 4, 4, f"{c_ev_outros:02d}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        
        self.merge_and_set(table_pernoite, 16, 0, 3, f"TOTAL EM OUTROS DESTINOS: {NUM_WORDS.get(total_outros, str(total_outros))}", bold=True)
        self.merge_and_set(table_pernoite, 16, 4, 4, f"{total_outros:02d}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        
        self.merge_and_set(table_pernoite, 17, 0, 3, f"TOTAL GERAL: {NUM_WORDS.get(total_geral, str(total_geral))}", bold=True)
        self.merge_and_set(table_pernoite, 17, 4, 4, f"{total_geral:02d}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        
        widths_p = [Inches(0.8), Inches(1.7), Inches(1.7), Inches(1.7), Inches(1.0)]
        for row in table_pernoite.rows:
            for idx, width in enumerate(widths_p):
                try: row.cells[idx].width = width
                except: pass

        # --- ASSINATURAS E ALTERAÇÃO ---
        table_sig = document.add_table(rows=2, cols=3)
        table_sig.alignment = WD_TABLE_ALIGNMENT.CENTER
        table_sig.style = 'Table Grid'
        table_sig.autofit = False
        widths_sig = [Inches(1.5), Inches(3.9), Inches(1.5)]
        
        c_sgt = table_sig.cell(0, 0)
        p_sgt = c_sgt.paragraphs[0]
        p_sgt.text = "Visto:"
        p_sgt.runs[0].font.size = Pt(8)
        p_sgt_line = c_sgt.add_paragraph()
        p_sgt_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sgt_line.paragraph_format.space_before = Pt(25)
        p_sgt_line.add_run("_________").font.size = Pt(8)
        run_sgt = p_sgt_line.add_run("\nSgt Dia")
        run_sgt.font.size = Pt(9)
        run_sgt.bold = True; run_sgt.underline = True
        
        c_mid_sig = table_sig.cell(0, 1)
        nome_sgte = self.current_state.get('nome_sgte', 'HEBERT CARLOS VIANA - 2° Sgt')
        p_date = c_mid_sig.paragraphs[0]
        p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_date = p_date.add_run(f"Quartel em Natal/RN, {d_str_yesterday}.")
        r_date.bold = True; r_date.font.size = Pt(10)
        p_name = c_mid_sig.add_paragraph()
        p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_name.paragraph_format.space_before = Pt(10)
        r_name = p_name.add_run(nome_sgte)
        r_name.bold = True; r_name.font.size = Pt(11)
        p_title = c_mid_sig.add_paragraph(f"Sargenteante da {u_nome.title()}")
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p_title.runs: r.font.size = Pt(9); r.underline = True
        
        c_of = table_sig.cell(0, 2)
        p_of = c_of.paragraphs[0]
        p_of.text = "Visto:"
        p_of.runs[0].font.size = Pt(8)
        p_of_line = c_of.add_paragraph()
        p_of_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_of_line.paragraph_format.space_before = Pt(25)
        p_of_line.add_run("_________").font.size = Pt(8)
        run_of = p_of_line.add_run("\nOf Dia")
        run_of.font.size = Pt(9)
        run_of.bold = True; run_of.underline = True

        c_alt = table_sig.cell(1, 0)
        c_alt.merge(table_sig.cell(1, 1)).merge(table_sig.cell(1, 2))
        p_alt = c_alt.paragraphs[0]
        run_alt = p_alt.add_run("Alteração: Com alteração (      ) Sem alteração (      )")
        run_alt.bold = True; run_alt.font.size = Pt(9)
        p_alt.paragraph_format.line_spacing = Pt(12)
        for _ in range(4):
            p_l = c_alt.add_paragraph()
            p_l.paragraph_format.line_spacing = Pt(12)
            p_l.paragraph_format.tab_stops.add_tab_stop(Inches(6.6), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.LINES)
            p_l.add_run("\t").font.size = Pt(9)

        for r_idx in range(2):
            for c_idx in range(3): table_sig.cell(r_idx, c_idx).width = widths_sig[c_idx]

        document.save(output_path)
        return output_path
